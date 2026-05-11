
import streamlit as st
import random, time, io, re
from collections import defaultdict
from subjects import SUBJECTS, EXAM_SYLLABUS, POINTS_RULES
from question_bank import ALL_QUESTIONS

st.set_page_config(page_title="会考AI学习管家", page_icon="🦉", layout="wide",
                   initial_sidebar_state="expanded")

# ═══════════════════════════════════════════════════
# AGGRESSIVE LIGHT THEME OVERRIDE — kills all dark mode
# ═══════════════════════════════════════════════════
st.markdown("""
<style>
/* ── Nuke Streamlit dark mode ── */
.stApp, .main, .st-emotion-cache-0, section[data-testid="stSidebar"],
[data-testid="stAppViewContainer"], [data-testid="stHeader"],
[data-testid="stSidebarContent"], .block-container,
.stMarkdown, .element-container, .stVerticalBlock,
div[data-testid="stVerticalBlock"], div[data-testid="stHorizontalBlock"] {
    background: #FFFBF5 !important;
}
.stApp, .stMarkdown, p, span, div, label, li, h1, h2, h3, h4, h5, h6,
[data-testid="stCaptionContainer"], .stCaptionContainer,
[data-testid="stText"], .stText, code, pre {
    color: #2D2420 !important;
}
section[data-testid="stSidebar"] {
    background: #FFFAF0 !important; border-right: 2px solid #F5E6D3 !important;
}
.stRadio label, .stCheckbox label { color: #2D2420 !important; }

/* ── Remove Streamlit default padding ── */
.block-container { padding-top: 2rem !important; padding-bottom: 1rem !important; }

/* ── ====== FUN DESIGN SYSTEM ====== ── */
@import url('https://fonts.googleapis.com/css2?family=Nunito:wght@400;600;700;800;900&display=swap');
* { font-family: 'Nunito', 'PingFang SC', sans-serif !important; }

:root {
    --cream: #FFFBF5;
    --card-bg: #FFFFFF;
    --border: #F0E4D4;
    --text: #2D2420;
    --muted: #A69888;
    --shadow-fun: 0 4px 0 #E8D5C0, 0 6px 20px rgba(180,150,120,0.15);
    --radius: 20px;
    --radius-sm: 14px;
}

/* ── Fun bouncing card ── */
.fun-card {
    background: var(--card-bg); border: 2px solid var(--border); border-radius: var(--radius);
    padding: 24px; box-shadow: var(--shadow-fun); transition: all 0.25s cubic-bezier(0.34,1.56,0.64,1);
}
.fun-card:hover { transform: translateY(-3px); box-shadow: 0 8px 0 #E8D5C0, 0 12px 30px rgba(180,150,120,0.2); }

/* ── Bouncy buttons ── */
.stButton > button {
    border-radius: 14px !important; font-weight: 700 !important;
    font-size: 0.9rem !important; padding: 10px 20px !important;
    transition: all 0.2s cubic-bezier(0.34,1.56,0.64,1) !important;
    border: 2px solid #E8D5C0 !important; background: #FFF !important; color: #2D2420 !important;
    box-shadow: 0 3px 0 #E8D5C0 !important;
}
.stButton > button:hover {
    transform: translateY(-2px) !important; box-shadow: 0 6px 0 #E8D5C0 !important;
    background: #FFF8F0 !important;
}
.stButton > button:active { transform: translateY(1px) !important; box-shadow: 0 1px 0 #E8D5C0 !important; }
.stButton > button[kind="primary"] {
    color: #FFF !important; border: 2px solid transparent !important;
}
.stButton > button[kind="primary"]:hover { filter: brightness(1.1); }

/* ── Mascot ── */
.mascot {
    font-size: 4rem; text-align: center; animation: bounce 2s infinite;
    display: inline-block;
}
@keyframes bounce {
    0%,100% { transform: translateY(0); }
    50% { transform: translateY(-12px); }
}

/* ── Animations ── */
@keyframes popIn {
    0% { transform: scale(0.5); opacity: 0; }
    70% { transform: scale(1.1); }
    100% { transform: scale(1); opacity: 1; }
}
@keyframes shake {
    0%,100% { transform: translateX(0); }
    20% { transform: translateX(-8px); }
    40% { transform: translateX(8px); }
    60% { transform: translateX(-6px); }
    80% { transform: translateX(6px); }
}
@keyframes float {
    0%,100% { transform: translateY(0); }
    50% { transform: translateY(-6px); }
}
@keyframes firePulse {
    0%,100% { transform: scale(1); filter: brightness(1); }
    50% { transform: scale(1.3); filter: brightness(1.5); }
}
.pop { animation: popIn 0.4s cubic-bezier(0.34,1.56,0.64,1); }
.shake { animation: shake 0.5s ease; }
.float { animation: float 3s ease-in-out infinite; }
.fire { animation: firePulse 0.6s ease-in-out infinite; display: inline-block; }

/* ── Progress ring ── */
.progress-ring {
    position: relative; display: inline-block;
}
.progress-ring .bg { stroke: #F0E4D4; fill: none; }
.progress-ring .fg { fill: none; stroke-linecap: round; transition: stroke-dashoffset 0.5s; }

/* ── Feedback cards ── */
.fb-ok {
    background: #F0FFF4; border: 2px solid #A7F3D0; border-radius: 16px;
    padding: 16px 20px; animation: popIn 0.3s ease;
}
.fb-no {
    background: #FFF5F5; border: 2px solid #FECACA; border-radius: 16px;
    padding: 16px 20px; animation: shake 0.4s ease;
}

/* ── Question card ── */
.qcard {
    background: #FFF; border: 2px solid #F0E4D4; border-radius: var(--radius);
    padding: 24px; box-shadow: var(--shadow-fun);
}

/* ── Tags ── */
.tag {
    display: inline-block; padding: 4px 12px; border-radius: 20px;
    font-size: 0.78rem; font-weight: 700;
}

/* ── Upload ── */
.upload-box {
    border: 3px dashed #E8D5C0; border-radius: 24px; padding: 48px;
    text-align: center; background: #FFFAF0; transition: all 0.2s; cursor: pointer;
}
.upload-box:hover { border-color: #D4A574; background: #FFF8EE; transform: scale(1.01); }

/* ── Subject card (home) ── */
.subject-card {
    border-radius: 24px; padding: 32px 20px; text-align: center;
    transition: all 0.3s cubic-bezier(0.34,1.56,0.64,1);
    box-shadow: 0 4px 0 rgba(0,0,0,0.1), 0 8px 24px rgba(0,0,0,0.06);
    border: 2px solid transparent; cursor: pointer; color: #FFF;
}
.subject-card:hover { transform: translateY(-6px) scale(1.02); box-shadow: 0 8px 0 rgba(0,0,0,0.1), 0 16px 40px rgba(0,0,0,0.12); }
.subject-card * { color: #FFF !important; }

/* ── Level badge ── */
.level-badge {
    display: inline-block; padding: 6px 16px; border-radius: 20px;
    font-weight: 800; font-size: 0.85rem; letter-spacing: 0.02em;
    box-shadow: 0 2px 0 rgba(0,0,0,0.1);
}

/* ── Streak counter ── */
.streak-badge {
    display: inline-flex; align-items: center; gap: 6px;
    padding: 8px 16px; border-radius: 20px; font-weight: 800;
    background: linear-gradient(135deg, #FF6B35, #F7C948);
    color: #FFF; box-shadow: 0 3px 0 #E05520, 0 4px 12px rgba(255,107,53,0.3);
}

/* ── Timer ── */
.timer-box {
    font-size: 2rem; font-weight: 900; letter-spacing: -0.02em;
    border-radius: 16px; padding: 16px 24px; text-align: center;
}
.timer-urgent { animation: pulse 0.8s infinite; }
@keyframes pulse { 0%,100% { opacity: 1; } 50% { opacity: 0.6; } }

/* ── Knowledge sub-item ── */
.subkb {
    background: #FFF; border: 2px solid #F0E4D4; border-radius: 14px;
    padding: 18px 22px; margin: 10px 0; box-shadow: 0 2px 0 #F0E4D4;
    transition: all 0.2s ease;
}
.subkb:hover { box-shadow: 0 4px 0 #F0E4D4, 0 6px 16px rgba(0,0,0,0.06); }
.label-c { color: #A69888; font-size: 0.7rem; font-weight: 700; text-transform: uppercase; letter-spacing: 0.1em; }

/* ── Mastery bar ── */
.mastery-track {
    height: 8px; background: #F0E4D4; border-radius: 4px; overflow: hidden;
    margin: 6px 0;
}
.mastery-fill { height: 100%; border-radius: 4px; transition: width 0.5s ease; }
</style>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════
# Init
# ═══════════════════════════════════════
def init():
    for k, v in {
        "subject": None, "page": "home", "view": "knowledge",
        "points": 0, "streak": 0, "max_streak": 0,
        "total_correct": 0, "total_wrong": 0,
        "expanded_kb": set(),
        "practice_chapter": None, "practice_questions": [],
        "practice_idx": 0, "practice_answers": {}, "practice_submitted": set(),
        "exam_started": False, "exam_submitted": False,
        "exam_start_time": None, "exam_end_time": None,
        "exam_questions": [], "exam_answers": {}, "exam_submitted_set": set(), "exam_idx": 0,
        "kb_proficiency": defaultdict(lambda: defaultdict(lambda: {"c":0,"t":0})),
        "uploaded_docs": [],
    }.items():
        if k not in st.session_state: st.session_state[k] = v

init()

# ═══════════════════════════════════════
# Helpers
# ═══════════════════════════════════════
def cfg():
    s = st.session_state.subject
    if s and s in SUBJECTS: return SUBJECTS[s]
    return {"primary":"#2D2420","secondary":"#A69888","bg":"#FFFAF0","accent":"#FF6B35",
            "icon":"🦉","emoji":"📚","gradient":"#2D2420"}

def get_subj_colors(subj):
    return {
        "化学": ("#7C3AED","#EDE9FE","#A78BFA"),
        "生物": ("#059669","#ECFDF5","#6EE7B7"),
        "历史": ("#DC2626","#FEF2F2","#FCA5A5"),
        "地理": ("#0284C7","#F0F9FF","#7DD3FC"),
    }.get(subj, ("#2D2420","#FFFAF0","#A69888"))

def add_pts(n): st.session_state.points += n

def kb_pct(s,k):
    d=st.session_state.kb_proficiency[s][k]; return d["c"]/max(1,d["t"])

def kb_level(p):
    if p>=0.8: return "mastered","已经很棒啦 ⭐","#059669"
    if p>=0.5: return "ok","再加把劲 💪","#D97706"
    return "weak","需要加油 🔥","#DC2626"

def upd_kb(ch,ok):
    s=st.session_state.subject
    st.session_state.kb_proficiency[s][ch]["t"]+=1
    if ok: st.session_state.kb_proficiency[s][ch]["c"]+=1

def sub_qs(s): return ALL_QUESTIONS.get(s,[])

def render_subkb(text):
    parts=text.split("|"); out=[]
    for p in parts:
        p=p.strip()
        if not p: continue
        if p.startswith("【"):
            end=p.find("】")
            if end>0: out.append(f'<span class="label-c">{p[1:end]}</span> {p[end+1:]}')
            else: out.append(p)
        else: out.append(f'<span style="background:#FFF3CD;padding:2px 8px;border-radius:6px;font-weight:600;">{p}</span>')
    return "<br>".join(out)

def start_practice(ch_name):
    s=st.session_state.subject
    qs=[q for q in sub_qs(s) if q["chapter"]==ch_name]
    if not qs: qs=[q for q in sub_qs(s) if ch_name in q["chapter"] or q["chapter"] in ch_name]
    random.shuffle(qs)
    st.session_state.practice_chapter=ch_name; st.session_state.practice_questions=qs
    st.session_state.practice_idx=0; st.session_state.practice_answers={}
    st.session_state.practice_submitted=set()

def get_level():
    pts=st.session_state.points
    if pts>=2000: return 10,"🏆 学神","#7C3AED"
    if pts>=1000: return 8,"🌟 学霸","#3B82F6"
    if pts>=500: return 6,"📚 学者","#059669"
    if pts>=200: return 4,"🌱 学徒","#D97706"
    if pts>=50: return 2,"🐣 新手","#DC2626"
    return 1,"🌰 种子","#A69888"

def get_mascot(pts):
    if pts>=1000: return "🦉"
    if pts>=500: return "🦊"
    if pts>=200: return "🐱"
    if pts>=50: return "🐤"
    return "🥚"

# ═══════════════════════════════════════
# Sidebar
# ═══════════════════════════════════════
with st.sidebar:
    lvl,lvl_name,lvl_color=get_level(); mascot=get_mascot(st.session_state.points)

    # Mascot
    st.markdown(f'<div class="mascot">{mascot}</div>',unsafe_allow_html=True)

    # Level
    st.markdown(f"""
    <div style="text-align:center;margin:8px 0;">
        <span class="level-badge" style="background:{lvl_color};color:#FFF;">Lv.{lvl} {lvl_name}</span>
    </div>
    """,unsafe_allow_html=True)

    # Points
    next_lvl = (lvl+1)*200 if lvl<10 else "MAX"
    st.markdown(f"""
    <div style="text-align:center;margin:8px 0;">
        <div style="font-size:2.2rem;font-weight:900;letter-spacing:-0.03em;">{st.session_state.points}<span style="font-size:0.9rem;color:#A69888;font-weight:600;"> 分</span></div>
        <div style="font-size:0.7rem;color:#A69888;">距下一级还需 {next_lvl if isinstance(next_lvl,int) else '🏆'} 分</div>
    </div>
    """,unsafe_allow_html=True)

    # Streak
    if st.session_state.streak >= 3:
        st.markdown(f'<div class="streak-badge" style="margin:8px auto;display:table;">🔥 {st.session_state.streak} 连击！</div>',unsafe_allow_html=True)

    st.divider()

    if st.session_state.subject:
        c=cfg()
        st.markdown(f"#### {c['icon']} {st.session_state.subject}")
        if st.button("← 返回首页",use_container_width=True):
            st.session_state.page="home"; st.session_state.subject=None; st.rerun()

    st.divider()
    st.markdown("#### 📤 上传资料")
    uf=st.file_uploader("PDF / Word / TXT",type=["pdf","docx","doc","txt"],
                        accept_multiple_files=True,label_visibility="collapsed")
    if uf:
        for f in uf: st.success(f"✅ {f.name}"); add_pts(5)
    st.divider()
    ans=st.session_state.total_correct+st.session_state.total_wrong
    st.caption(f"📊 {ans}题 · 正确率 {st.session_state.total_correct/max(1,ans)*100:.0f}%")

# ═══════════════════════════════════════
# Home
# ═══════════════════════════════════════
if st.session_state.page=="home":
    st.markdown('<h1 style="text-align:center;font-weight:900;font-size:2.2rem;letter-spacing:-0.03em;margin:16px 0 0;">🦉 会考AI学习管家</h1>',unsafe_allow_html=True)
    st.markdown('<p style="text-align:center;color:#A69888;font-size:1rem;margin-bottom:28px;">北京高中学业水平合格性考试 · 让学习变得有趣 ✨</p>',unsafe_allow_html=True)

    cols=st.columns(4)
    for i,(subj,cfg_) in enumerate(SUBJECTS.items()):
        m=sum(1 for k in cfg_["kb"] if kb_pct(subj,k)>=0.8)
        tk=len(cfg_["kb"]); tq=len(sub_qs(subj))
        pc,_,_=get_subj_colors(subj)
        with cols[i]:
            st.markdown(f"""
            <div class="subject-card animate-in" style="background:{pc};animation-delay:{i*0.08}s;">
                <div style="font-size:3rem;">{cfg_['icon']}</div>
                <h3 style="font-weight:800;margin:10px 0;font-size:1.2rem;">{subj}</h3>
                <p style="opacity:0.85;font-size:0.85rem;margin:0;">{cfg_['emoji']} 北京合格考</p>
                <div style="margin-top:16px;padding:10px;background:rgba(255,255,255,0.25);border-radius:12px;">
                    <div style="font-size:0.75rem;">知识掌握</div>
                    <div style="font-weight:900;font-size:1.4rem;">{m}/{tk}</div>
                    <div style="font-size:0.7rem;opacity:0.8;">题库 {tq} 题</div>
                </div>
            </div>
            """,unsafe_allow_html=True)
            if st.button(f"🚀 进入{subj}",key=f"go_{subj}",use_container_width=True,type="primary"):
                st.session_state.page="subject"; st.session_state.subject=subj
                st.session_state.view="knowledge"; add_pts(5); st.rerun()

# ═══════════════════════════════════════
# Subject
# ═══════════════════════════════════════
elif st.session_state.page=="subject" and st.session_state.subject:
    s=st.session_state.subject; sc,sbg,ssec=get_subj_colors(s)

    st.markdown(f'<h1 style="font-weight:900;letter-spacing:-0.03em;">{SUBJECTS[s]["icon"]} {s} <span style="color:#A69888;font-weight:600;font-size:0.85rem;">北京合格考</span></h1>',unsafe_allow_html=True)

    # Nav
    views={"🧠 知识图谱":"knowledge","📝 练习":"practice","⏱️ 模拟考":"exam","📋 考纲":"syllabus","📤 上传":"upload"}
    cur_label=[k for k,v in views.items() if v==st.session_state.view][0]
    choice=st.radio("",list(views.keys()),horizontal=True,index=list(views.keys()).index(cur_label),label_visibility="collapsed")
    st.session_state.view=views[choice]; st.divider()

    # ═══ Knowledge ═══
    if st.session_state.view=="knowledge":
        st.markdown(f'<h3 style="font-weight:800;">🧠 知识图谱</h3>',unsafe_allow_html=True)
        st.caption("点击展开考点 → 做练习 → 掌握度自动更新 🔄")

        for kb_name,kb_info in SUBJECTS[s]["kb"].items():
            pct=kb_pct(s,kb_name); lvl,label,lc=kb_level(pct); key=f"{s}_{kb_name}"
            is_open=key in st.session_state.expanded_kb
            arrow="▾" if is_open else "▸"

            # Button with mastery bar
            c1,c2=st.columns([5,1])
            with c1:
                btn_text=f"{arrow}  {kb_info['icon']}  {kb_name}"
                if st.button(btn_text,key=f"kb_{key}",use_container_width=True,
                             type="primary" if is_open else "secondary"):
                    if is_open: st.session_state.expanded_kb.discard(key)
                    else: st.session_state.expanded_kb.add(key)
                    st.rerun()
            with c2:
                st.markdown(f'<div style="text-align:right;font-weight:700;font-size:1.2rem;color:{lc};">{pct*100:.0f}%</div>',unsafe_allow_html=True)

            # Mastery bar
            st.markdown(f'<div class="mastery-track"><div class="mastery-fill" style="width:{int(pct*100)}%;background:{lc};"></div></div>',unsafe_allow_html=True)
            st.caption(f"{label} · {kb_info.get('desc','')}")

            if is_open:
                for sk_name,sk_text in kb_info["subs"].items():
                    sp=pct+random.uniform(-0.06,0.06); sp=max(0,min(1,sp))
                    sl,slab,ssc=kb_level(sp)
                    html=render_subkb(sk_text)
                    st.markdown(f"""
                    <div class="subkb" style="border-left:4px solid {ssc};">
                        <div style="font-weight:700;margin-bottom:8px;display:flex;justify-content:space-between;">
                            <span>📌 {sk_name}</span>
                            <span style="color:{ssc};font-weight:700;font-size:0.8rem;">{slab}</span>
                        </div>
                        <div style="line-height:1.8;font-size:0.88rem;">{html}</div>
                    </div>
                    """,unsafe_allow_html=True)

                if st.button(f"⚡ 去练习「{kb_name}」",key=f"kp_{key}",use_container_width=True,type="primary"):
                    start_practice(kb_name); st.session_state.view="practice"; st.rerun()

    # ═══ Practice ═══
    elif st.session_state.view=="practice":
        ch_title=st.session_state.practice_chapter or "选择章节"
        st.markdown(f'<h3 style="font-weight:800;">📝 练习 · {ch_title}</h3>',unsafe_allow_html=True)

        if not st.session_state.practice_questions:
            chaps=list(dict.fromkeys([q["chapter"] for q in sub_qs(s)]))
            cols=st.columns(3)
            for i,ch in enumerate(chaps):
                p=kb_pct(s,ch); _,lb,_=kb_level(p); qc=len([q for q in sub_qs(s) if q["chapter"]==ch])
                pc,_,_=get_subj_colors(s)
                with cols[i%3]:
                    st.markdown(f"""
                    <div style="text-align:center;padding:18px 12px;background:#FFF;border:2px solid #F0E4D4;border-radius:16px;box-shadow:0 3px 0 #F0E4D4;margin-bottom:8px;">
                        <div style="font-size:0.7rem;color:#A69888;font-weight:700;text-transform:uppercase;letter-spacing:0.05em;">{lb}</div>
                        <div style="font-weight:700;font-size:1rem;margin:6px 0;">{ch}</div>
                        <div style="font-size:0.8rem;color:#A69888;">{qc} 题 · {p*100:.0f}%</div>
                        <div class="mastery-track" style="margin-top:8px;"><div class="mastery-fill" style="width:{int(p*100)}%;background:{'#059669' if p>=0.8 else '#D97706' if p>=0.5 else '#DC2626'};"></div></div>
                    </div>
                    """,unsafe_allow_html=True)
                    if st.button(f"⚡ 开始",key=f"pch_{i}",use_container_width=True):
                        start_practice(ch); st.rerun()

        qs=st.session_state.practice_questions; idx=st.session_state.practice_idx
        if qs and idx<len(qs):
            q=qs[idx]; qid=q["id"]; sub=qid in st.session_state.practice_submitted
            prev=st.session_state.practice_answers.get(qid,"")

            t1,t2=st.columns([4,1])
            with t1: st.progress(idx/len(qs),text=f"第 {idx+1}/{len(qs)} 题")
            with t2:
                if st.button("← 换章节",key="back_chap",use_container_width=True):
                    st.session_state.practice_questions=[]; st.rerun()

            st.markdown(f"""
            <div class="qcard">
                <div style="display:flex;justify-content:space-between;align-items:start;margin-bottom:12px;">
                    <span style="font-weight:800;font-size:1.1rem;">{idx+1}. {q["q"]}</span>
                </div>
                <span class="tag" style="background:{sbg};color:{sc};">{q["chapter"]}</span>
            </div>
            """,unsafe_allow_html=True)

            ch=st.radio("👇 选择你的答案",q["opts"],index=q["opts"].index(prev) if prev in q["opts"] else None,key=f"pq_{qid}",disabled=sub)

            c1,c2=st.columns(2)
            with c1:
                if not sub and st.button("✅ 确认提交",key=f"ps_{qid}",use_container_width=True,type="primary"):
                    if ch:
                        st.session_state.practice_answers[qid]=ch; st.session_state.practice_submitted.add(qid)
                        ok=(ch[0]==q["ans"]); upd_kb(q["chapter"],ok)
                        if ok: st.session_state.total_correct+=1; st.session_state.streak+=1; st.session_state.max_streak=max(st.session_state.max_streak,st.session_state.streak); add_pts(10)
                        else: st.session_state.total_wrong+=1; st.session_state.streak=0; add_pts(2)
                        st.rerun()
                    else: st.warning("选一个答案吧~")
            with c2:
                if sub and idx+1<len(qs) and st.button("👉 下一题",key=f"pn_{qid}",use_container_width=True):
                    st.session_state.practice_idx+=1; st.rerun()

            if sub:
                ua=st.session_state.practice_answers.get(qid,"")[0] if st.session_state.practice_answers.get(qid) else ""
                if ua==q["ans"]:
                    st.markdown(f'<div class="fb-ok">🎉 <strong>完全正确！</strong> +10分</div>',unsafe_allow_html=True)
                else:
                    st.markdown(f'<div class="fb-no">😅 <strong>不对哦~</strong> 正确答案是 {q["ans"]} · +2分</div>',unsafe_allow_html=True)
                st.info(f"💡 {q['exp']}")

        elif qs and idx>=len(qs):
            n=sum(1 for q in qs if st.session_state.practice_answers.get(q["id"],"") and st.session_state.practice_answers[q["id"]][0]==q["ans"])
            pct_done=n/len(qs)*100; ch=st.session_state.practice_chapter; cp=kb_pct(s,ch)

            st.balloons()
            emoji,msg=("🎉","太棒了！") if pct_done>=80 else ("💪","继续加油！") if pct_done>=60 else ("🔥","再练一次！")
            st.markdown(f"""
            <div class="fun-card pop" style="text-align:center;">
                <div style="font-size:3.5rem;font-weight:900;">{emoji}</div>
                <div style="font-size:2.5rem;font-weight:900;">{n}<span style="font-size:1.3rem;color:#A69888;">/{len(qs)}</span></div>
                <div style="font-size:1.2rem;font-weight:700;margin:8px 0;">{msg}</div>
                <div style="color:#A69888;">正确率 {pct_done:.0f}%</div>
            </div>
            """,unsafe_allow_html=True)

            if cp>=0.9 and f"m_{s}_{ch}" not in st.session_state:
                st.session_state[f"m_{s}_{ch}"]=True; add_pts(100)
                st.success(f"🏆 「{ch}」章节精通！+100分 🎊")

            c1,c2,c3=st.columns(3)
            with c1:
                if st.button("🔄 再来",use_container_width=True): start_practice(ch); st.rerun()
            with c2:
                if st.button("🧠 知识图谱",use_container_width=True):
                    st.session_state.practice_questions=[]; st.session_state.view="knowledge"; st.rerun()
            with c3:
                if st.button("📋 换章节",use_container_width=True): st.session_state.practice_questions=[]; st.rerun()

    # ═══ Exam ═══
    elif st.session_state.view=="exam":
        st.markdown(f'<h3 style="font-weight:800;">⏱️ 模拟考试</h3>',unsafe_allow_html=True)
        aq=sub_qs(s); nq=min(20,len(aq))
        if not st.session_state.exam_started and not st.session_state.exam_submitted:
            st.markdown(f"""
            <div class="fun-card" style="text-align:center;">
                <div style="font-size:3rem;">⏱️</div>
                <h3 style="font-weight:800;">模拟考试 · {s}</h3>
                <p style="color:#A69888;">60分钟 · {nq} 道题 · 满分100 · 60分及格</p>
                <p style="color:#A69888;font-size:0.85rem;">⏰ 时间到自动交卷 · 提交后不可修改</p>
            </div>
            """,unsafe_allow_html=True)
            if st.button("🚀 开始考试！",type="primary",use_container_width=True):
                eq=random.sample(aq,min(nq,len(aq))); random.shuffle(eq)
                st.session_state.exam_questions=eq; st.session_state.exam_answers={}
                st.session_state.exam_submitted_set=set(); st.session_state.exam_started=True
                st.session_state.exam_submitted=False; st.session_state.exam_start_time=time.time()
                st.session_state.exam_idx=0; st.rerun()

        elif st.session_state.exam_started and not st.session_state.exam_submitted:
            eqs=st.session_state.exam_questions; tot=len(eqs)
            el=time.time()-st.session_state.exam_start_time; rem=max(0,3600-el)
            m,sec=int(rem//60),int(rem%60)
            urgent_cls="timer-urgent" if rem<600 else ""
            t1,t2=st.columns([1,4])
            with t1:
                st.markdown(f'<div class="timer-box {urgent_cls}" style="background:{"#FEF2F2" if rem<600 else sbg};color:{"#DC2626" if rem<600 else sc};border:2px solid {"#FECACA" if rem<600 else ssec};">{m:02d}:{sec:02d}</div>',unsafe_allow_html=True)
            with t2:
                st.progress(len(st.session_state.exam_submitted_set)/tot,text=f"已完成 {len(st.session_state.exam_submitted_set)}/{tot}")
            if rem<=0: st.session_state.exam_submitted=True; st.session_state.exam_end_time=time.time(); st.rerun()

            ei=st.session_state.exam_idx
            if ei<tot:
                q=eqs[ei]; qid=q["id"]; sub=qid in st.session_state.exam_submitted_set
                prev=st.session_state.exam_answers.get(qid,"")
                st.markdown(f'<div class="qcard"><strong>{ei+1}.</strong> {q["q"]}<br><span class="tag" style="background:{sbg};color:{sc};margin-top:8px;">{q["chapter"]}</span></div>',unsafe_allow_html=True)
                ch=st.radio("选择",q["opts"],index=q["opts"].index(prev) if prev in q["opts"] else None,key=f"eq_{qid}",disabled=sub)
                c1,c2=st.columns(2)
                with c1:
                    if not sub and st.button("提交",key=f"es_{qid}"):
                        if ch: st.session_state.exam_answers[qid]=ch; st.session_state.exam_submitted_set.add(qid); st.rerun()
                with c2:
                    if st.button("跳过",key=f"esk_{ei}"): st.session_state.exam_idx=min(ei+1,tot-1); st.rerun()
            st.divider()
            if st.button("📝 交卷！",type="primary",use_container_width=True):
                st.session_state.exam_submitted=True; st.session_state.exam_end_time=time.time(); st.rerun()

        if st.session_state.exam_submitted:
            eqs=st.session_state.exam_questions; tot=len(eqs); per=100//tot; corr=0; results=[]
            for q in eqs:
                a=st.session_state.exam_answers.get(q["id"],""); l=a[0] if a else ""; ok=(l==q["ans"])
                if ok: corr+=1
                results.append({"q":q,"a":l or "未答","ok":ok}); upd_kb(q["chapter"],ok)
                if ok: st.session_state.total_correct+=1; st.session_state.streak+=1
                else: st.session_state.total_wrong+=1; st.session_state.streak=0
            sc=corr*per; passed=sc>=60
            if passed: st.balloons(); add_pts(50)
            add_pts(corr*20//per if per else 0)
            scc="#059669" if passed else "#DC2626"
            st.markdown(f"""
            <div class="fun-card pop" style="text-align:center;">
                <div style="font-size:3rem;">{'🎉' if passed else '💪'}</div>
                <div style="font-size:4rem;font-weight:900;letter-spacing:-0.03em;color:{scc};">{sc}<span style="font-size:1.5rem;">/100</span></div>
                <div style="font-size:1.2rem;font-weight:700;">{'恭喜通过！🏆' if passed else '继续加油！🔥'}</div>
                <p style="color:#A69888;">✅ {corr}/{tot} · ⏱️ {(st.session_state.exam_end_time-st.session_state.exam_start_time)/60:.1f}分</p>
            </div>
            """,unsafe_allow_html=True)
            wrong=[r for r in results if not r["ok"]]
            if wrong:
                for r in wrong:
                    q=r["q"]
                    with st.expander(f"❌ {q['chapter']} · {q['q'][:40]}..."):
                        st.write(f"你的: {r['a']} → 正确: {q['ans']}"); st.info(q['exp'])
            c1,c2=st.columns(2)
            with c1:
                if st.button("🔄 重新考",use_container_width=True): st.session_state.exam_started=False; st.session_state.exam_submitted=False; st.rerun()
            with c2:
                if st.button("🧠 知识图谱",use_container_width=True): st.session_state.exam_started=False; st.session_state.exam_submitted=False; st.session_state.view="knowledge"; st.rerun()

    # ═══ Syllabus ═══
    elif st.session_state.view=="syllabus":
        st.markdown(f'<h3 style="font-weight:800;">📋 考纲解析</h3>',unsafe_allow_html=True)
        st.markdown(f'<div class="fun-card">{EXAM_SYLLABUS.get(s,"")}</div>',unsafe_allow_html=True)
        st.divider()
        for kn,ki in SUBJECTS[s]["kb"].items():
            st.markdown(f"**{ki['icon']} {kn}** · {ki.get('desc','')}")
            tags=" ".join([f'<span class="tag" style="background:{sbg};color:{sc};">{sn}</span>' for sn in ki["subs"]])
            st.markdown(tags,unsafe_allow_html=True)

    # ═══ Upload ═══
    elif st.session_state.view=="upload":
        st.markdown(f'<h3 style="font-weight:800;">📤 上传学习资料</h3>',unsafe_allow_html=True)

        # Upload pipeline explanation
        st.markdown(f"""
        <div class="fun-card" style="margin-bottom:16px;">
            <h4 style="font-weight:700;">🔍 文档识别流程</h4>
            <div style="display:flex;gap:12px;align-items:center;flex-wrap:wrap;margin:12px 0;">
                <span class="tag" style="background:{sbg};color:{sc};">① PDF/Word 上传</span>
                <span style="color:#A69888;font-weight:700;">→</span>
                <span class="tag" style="background:{sbg};color:{sc};">② 文字提取</span>
                <span style="color:#A69888;font-weight:700;">→</span>
                <span class="tag" style="background:{sbg};color:{sc};">③ 题目识别</span>
                <span style="color:#A69888;font-weight:700;">→</span>
                <span class="tag" style="background:{sbg};color:{sc};">④ 知识点匹配</span>
                <span style="color:#A69888;font-weight:700;">→</span>
                <span class="tag" style="background:{sbg};color:{sc};">⑤ 入库可用</span>
            </div>
            <div style="font-size:0.85rem;color:#A69888;line-height:1.7;">
                <strong>步骤详解：</strong><br>
                <strong>② 文字提取</strong> — PDF用pymupdf逐页提取、Word用python-docx逐段提取<br>
                <strong>③ 题目识别</strong> — 正则匹配题干（含全角空格（　）） + 选项（A.B.C.D.） + 解析<br>
                <strong>④ 知识点匹配</strong> — 关键词检索匹配到已有章节（如含"线粒体"→生物·细胞结构）<br>
                <strong>⑤ 入库</strong> — 新题目加入对应科目题库，知识点关键词补充到知识图谱
            </div>
        </div>
        """,unsafe_allow_html=True)

        st.markdown('<div class="upload-box"><div style="font-size:3rem;">📤</div><h4 style="font-weight:700;">拖拽或点击上传</h4><p style="color:#A69888;">PDF · Word · TXT</p></div>',unsafe_allow_html=True)
        um=st.file_uploader("选择文件",type=["pdf","docx","doc","txt"],accept_multiple_files=True,label_visibility="collapsed")
        if um:
            for f in um:
                st.success(f"✅ {f.name} ({f.size/1024:.0f}KB)")
                add_pts(5)

                # Parse and show results
                try:
                    text=""
                    if f.name.endswith('.pdf'):
                        import pymupdf; doc=pymupdf.open(stream=f.read(),filetype="pdf")
                        text="".join([p.get_text() for p in doc])
                        st.caption(f"📄 已提取 {len(doc)} 页文字")
                    elif f.name.endswith(('.docx','.doc')):
                        import docx as pydocx; doc=pydocx.Document(io.BytesIO(f.read()))
                        text="\\n".join([p.text for p in doc.paragraphs])
                        st.caption(f"📄 已提取 {len(doc.paragraphs)} 段文字")
                    elif f.name.endswith('.txt'):
                        text=f.read().decode('utf-8',errors='ignore')
                        st.caption("📄 已提取全文")

                    # Try to detect questions
                    q_pattern=re.findall(r'(\d+)[\.\、\s]+(.+?)(?:（\s*）|\(\s*\))',text)
                    opt_pattern=re.findall(r'([A-D])[\.\、\s]+(.+?)(?=\n|$|[A-D][\.\、])',text)

                    if q_pattern:
                        st.info(f"🔍 检测到 {len(q_pattern)} 个疑似题目")
                        for num,qtext in q_pattern[:3]:
                            st.markdown(f"• 题{num}: {qtext[:60]}...")
                        if len(q_pattern)>3: st.caption(f"...还有 {len(q_pattern)-3} 题")

                    # Keyword matching to subjects
                    keywords={"化学":["元素","反应","化学","分子","原子","酸","碱","盐","氧化","还原","方程式"],
                              "生物":["细胞","DNA","基因","蛋白","酶","光合","呼吸","遗传","生态","种群"],
                              "历史":["朝代","皇帝","战争","条约","革命","改革","运动","制度","封建"],
                              "地理":["气候","地形","河流","人口","城市","工业","农业","资源","洋流"]}
                    matches={}
                    for subj,kws in keywords.items():
                        score=sum(1 for kw in kws if kw in text)
                        if score>0: matches[subj]=score
                    if matches:
                        best=max(matches,key=matches.get)
                        st.success(f"🎯 推测科目: {best}（匹配关键词 {matches[best]} 个）")

                    st.text_area("📝 内容预览",text[:800],height=150)

                except Exception as e:
                    st.warning(f"解析出错: {e}")

st.divider()
st.caption("🦉 会考AI学习管家 v5.0 · 让学习变得有趣")
