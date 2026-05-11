"""文档处理模块 — 提取题目、匹配科目章节"""
import re, io, json, os
from pathlib import Path
from subjects import SUBJECTS

# 科目关键词库（用于自动匹配）
SUBJECT_KEYWORDS = {
    "化学": ["元素","反应","化学","分子","原子","酸","碱","盐","氧化","还原","方程式",
             "钠","铝","铁","氯","硫","氮","硅","周期","电解","原电池","摩尔","浓度",
             "离子","共价","催化剂","平衡","pH","有机","烷烃","烯烃","苯","乙醇"],
    "生物": ["细胞","DNA","RNA","基因","蛋白","酶","光合","呼吸","遗传","生态","种群",
             "群落","染色体","减数","有丝","转录","翻译","突变","进化","免疫","激素",
             "神经","反射","线粒体","叶绿体","核糖体","ATP","葡萄糖","氨基酸"],
    "历史": ["朝代","皇帝","战争","条约","革命","改革","运动","制度","封建","帝国",
             "秦","汉","唐","宋","元","明","清","民国","共和国","鸦片","辛亥革命",
             "五四","抗日","解放","改革开放","启蒙","工业革命","新航路","罗马"],
    "地理": ["气候","地形","河流","人口","城市","工业","农业","资源","洋流","经线",
             "纬线","等高线","季风","地中海","赤道","北极","南极","板块","地震",
             "火山","降水","温度","气压","风带","秦岭","淮河","长江","黄河"],
}

# 章节关键词映射
CHAPTER_KEYWORDS = {
    "化学": {
        "物质的量": ["摩尔","NA","阿伏伽德罗","浓度","mol","标况","气体摩尔体积"],
        "离子反应": ["电解质","电离","离子方程式","离子共存","沉淀"],
        "氧化还原反应": ["氧化","还原","化合价","氧化剂","还原剂","电子转移"],
        "金属及其化合物": ["钠","铝","铁","金属","合金","两性"],
        "非金属及其化合物": ["氯","硫","氮","硅","非金属","酸雨","硝酸"],
        "元素周期律": ["周期","族","原子半径","金属性","非金属性","化学键"],
        "化学反应与能量": ["速率","平衡","勒夏特列","放热","吸热","原电池","电解"],
        "有机化学基础": ["甲烷","乙烯","苯","乙醇","乙酸","酯化","取代","加成","官能团"],
    },
    "生物": {
        "细胞的分子组成": ["蛋白质","核酸","糖类","脂质","水","无机盐","氨基酸","核苷酸"],
        "细胞的结构": ["细胞膜","细胞器","线粒体","叶绿体","核糖体","细胞核","原核"],
        "细胞的代谢": ["酶","ATP","光合作用","呼吸作用","有氧","无氧"],
        "遗传的基本规律": ["分离定律","自由组合","孟德尔","伴性","显性","隐性","基因型"],
        "变异与进化": ["突变","基因重组","染色体变异","进化","自然选择","基因频率"],
        "稳态与环境": ["内环境","神经","体液","免疫","种群","群落","生态系统","食物链"],
    },
    "历史": {
        "中国古代政治": ["分封","宗法","郡县","三公","三省","科举","内阁","军机处"],
        "中国古代经济": ["农业","手工业","商业","丝绸之路","交子","商帮","井田制"],
        "近代中国(1840-1919)": ["鸦片","太平天国","洋务","甲午","戊戌","义和团","辛亥革命","条约"],
        "新民主主义革命": ["五四","中共","北伐","长征","抗日","解放战争"],
        "现代中国": ["建国","土改","一五","三大改造","改革开放","市场经济"],
        "世界史": ["希腊","罗马","文艺复兴","启蒙","工业革命","美国独立","法国大革命"],
    },
    "地理": {
        "地球与地图": ["经纬","本初","赤道","等高线","比例尺","方向"],
        "自然地理": ["大气","气压","风","气候","水循环","洋流","地貌","板块"],
        "人文地理": ["人口","城市化","农业","工业","交通","商业"],
        "区域地理": ["区域","可持续发展","荒漠化","水土流失","南水北调"],
        "中国地理": ["秦岭","淮河","长江","黄河","北方","南方","西北","青藏"],
    },
}

def extract_text(file_obj) -> tuple[str, str]:
    """从上传文件提取文本。返回 (文本, 方法描述)"""
    name = file_obj.name.lower()
    if name.endswith('.pdf'):
        import pymupdf
        doc = pymupdf.open(stream=file_obj.getvalue(), filetype="pdf")
        text = "".join([p.get_text() for p in doc])
        return text, f"pymupdf · {len(doc)}页"
    elif name.endswith(('.docx', '.doc')):
        import docx as pydocx
        doc = pydocx.Document(io.BytesIO(file_obj.getvalue()))
        text = "\n".join([p.text for p in doc.paragraphs])
        return text, f"python-docx · {len(doc.paragraphs)}段"
    elif name.endswith('.txt'):
        text = file_obj.getvalue().decode('utf-8', errors='ignore')
        return text, "txt直接读取"
    return "", "不支持格式"

def detect_subject(text: str) -> str:
    """根据文本内容自动检测科目"""
    scores = {}
    for subj, keywords in SUBJECT_KEYWORDS.items():
        score = sum(1 for kw in keywords if kw in text)
        if score > 0: scores[subj] = score
    if not scores: return None
    return max(scores, key=scores.get)

def detect_chapter(text: str, subject: str) -> str:
    """在给定科目下检测最匹配的章节"""
    if subject not in CHAPTER_KEYWORDS: return None
    scores = {}
    for ch, keywords in CHAPTER_KEYWORDS[subject].items():
        score = sum(1 for kw in keywords if kw in text)
        if score > 0: scores[ch] = score
    if not scores: return None
    return max(scores, key=scores.get)

def parse_questions(text: str) -> list[dict]:
    """从文本中解析题目。返回题目列表"""
    questions = []
    lines = text.split('\n')

    current_q = None
    for line in lines:
        line = line.strip()
        if not line: continue

        # 匹配题干：数字开头 + 中文内容 + 可能含（　）
        q_match = re.match(r'^(\d+)[\.\、\s）\)]\s*(.+?)(?:（\s*）|\(\s*\))\s*$', line)
        if q_match:
            if current_q:
                questions.append(current_q)
            current_q = {"q": q_match.group(2).strip(), "opts": [], "ans": "", "exp": ""}
            continue

        # 匹配选项 A. B. C. D.
        opt_match = re.match(r'^([A-D])[\.\、\s）\)]\s*(.+)', line)
        if opt_match and current_q is not None:
            current_q["opts"].append(f"{opt_match.group(1)}. {opt_match.group(2).strip()}")
            continue

        # 匹配答案
        ans_match = re.match(r'^(?:答案|正确答案)[：:]\s*([A-D])', line)
        if ans_match and current_q is not None:
            current_q["ans"] = ans_match.group(1)
            continue

        # 匹配解析
        exp_match = re.match(r'^(?:解析|说明)[：:]\s*(.+)', line)
        if exp_match and current_q is not None:
            current_q["exp"] = exp_match.group(1)
            continue

    if current_q and current_q.get("opts"):
        questions.append(current_q)

    # 过滤：至少要有题干、4个选项、答案
    valid = []
    for q in questions:
        if q.get("q") and len(q.get("opts",[])) >= 2:
            valid.append(q)
    return valid

def process_upload(username: str, file_obj) -> dict:
    """完整处理上传文件。返回处理结果摘要"""
    result = {
        "filename": file_obj.name,
        "size_kb": len(file_obj.getvalue()) / 1024,
        "text_method": "",
        "text_length": 0,
        "detected_subject": None,
        "detected_chapter": None,
        "questions_found": 0,
        "questions_saved": 0,
        "storage_path": "",
    }

    # 1. 提取文本
    text, method = extract_text(file_obj)
    result["text_method"] = method
    result["text_length"] = len(text)

    if not text:
        result["error"] = "无法提取文本"
        return result

    # 2. 检测科目
    subj = detect_subject(text)
    result["detected_subject"] = subj

    # 3. 检测章节
    if subj:
        ch = detect_chapter(text, subj)
        result["detected_chapter"] = ch

    # 4. 解析题目
    questions = parse_questions(text)
    result["questions_found"] = len(questions)

    # 5. 补充科目和章节信息
    for q in questions:
        q["subject"] = subj or "未知"
        q["chapter"] = ch or "未分类"
        q["source"] = file_obj.name
        q["id"] = f"up_{username}_{file_obj.name[:10]}_{len(q.get('opts',[]))}"

    # 6. 保存原始文件和提取结果
    from persistence import save_uploaded_file, save_extracted_questions
    result["storage_path"] = save_uploaded_file(username, file_obj)
    if questions:
        save_extracted_questions(username, file_obj.name, subj or "未知", questions)
        result["questions_saved"] = len(questions)

    return result
